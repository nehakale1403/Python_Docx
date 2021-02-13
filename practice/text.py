import docx
import os

doc = docx.Document()
doc.add_heading('Test Doc', 0)
doc.add_heading("Heading Level1", 1)

paragraph = doc.add_paragraph(input("Enter the inputs: "))
paragraph.add_run("This document is created using ")
paragraph.add_run("Python").bold=True

newparagraph = doc.add_paragraph()
newparagraph.add_run("This line is in italics!").italic=True
doc.save("test.docx")
os.system("start test.docx")